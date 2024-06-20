from docx import Document
from docx.shared import Pt, RGBColor
import json

def docx():

    def set_font(paragraph,op='',size=None):
        run = paragraph.runs[0]
        font = run.font
        font.size = Pt(size)
        if '√' in op:
            font.color.rgb = RGBColor(255, 0, 0)
        else:
            font.color.rgb = RGBColor(0, 0, 0)
    
    def write_no_save():
        document = Document()
        for i in dict.keys():
            heading = document.add_heading(i, 0)
            set_font(heading,size=11)

            for op in dict[i]:
                p = document.add_paragraph(op)
                set_font(p,op,size=11)
        return document
    
    json1 = open('E:/桌面/习概/习概预题库1.json','r',encoding='utf-8')
    dict = json.load(json1)
    document1 = write_no_save()
    document1.save('E:/桌面/习概/习概预题库1.docx')

    json2 = open('E:/桌面/习概/习概预题库2.json','r',encoding='utf-8')
    dict = json.load(json2)
    document2 = write_no_save()
    document2.save('E:/桌面/习概/习概预题库2.docx')
    
    
if __name__=='__main__':
    docx()