from docx import Document
from docx.shared import Pt, RGBColor
import json

def docx():
    def set_font(paragraph,op='',size=None):
        run = paragraph.runs[0]
        font = run.font
        font.size = Pt(size)
        if '√' in op:
            font.color.rgb = RGBColor(255, 0, 0)
        else:
            font.color.rgb = RGBColor(0, 0, 0)

    json1 = open('E:/桌面/毛概/毛概预题库.json','r',encoding='utf-8')
    dict = json.load(json1)
    document = Document()
    for i in dict.keys():
        heading = document.add_heading(i, 0)
        set_font(heading,size=11)

        for op in dict[i]:
            p = document.add_paragraph(op)
            set_font(p,op,size=11)

    document.save('E:/桌面/毛概/毛概预题库.docx')

if __name__=='__main__':
    docx()